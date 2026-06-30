from docx import Document
import os


def generate_document(user_request: str, content: str):
    """
    Generate a professional Word document.
    """

    os.makedirs("output", exist_ok=True)

    doc = Document()

    doc.add_heading("AI Generated Document", level=1)

    doc.add_heading("User Request", level=2)
    doc.add_paragraph(user_request)

    doc.add_heading("Generated Content", level=2)
    doc.add_paragraph(content)

    filename = "output/generated_document.docx"

    doc.save(filename)

    return filename